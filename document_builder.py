"""
Document builder — generates Word (.docx) and PDF files for resumes and cover letters.

Resume layout is hard-specced from Yang's actual PDFs. Haiku provides the tailored
text content; this module handles all formatting. Layout never changes.

Two resume templates:
  - Lending (Template A): black headings, horizontal rules between sections, bullet points
  - Risk/Controls (Template B): blue (#2E74B5) headings, name underline, dash bullets

Template is chosen automatically based on job type.

PDF generation: the Word document is built first, then converted to PDF via LibreOffice
headless. Both files are therefore byte-for-byte identical in layout.
LibreOffice is declared in packages.txt so Streamlit Cloud installs it automatically.
fonts-crosextra-carlito provides a Calibri-compatible font so the PDF matches the .docx.
"""

import io
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Literal, Optional
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import CLAUDE_MODEL, CONTROL_RISK_SIGNALS

_client: Optional[anthropic.Anthropic] = None
_PROFILE_PATH        = Path(__file__).parent / "data" / "candidate_profile.txt"
_RESUME_LENDING_PATH = Path(__file__).parent / "data" / "resume_general.txt"
_RESUME_RISK_PATH    = Path(__file__).parent / "data" / "resume_control_risk.txt"

BLUE      = RGBColor(0x2E, 0x74, 0xB5)
BLACK     = RGBColor(0x00, 0x00, 0x00)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_client() -> anthropic.Anthropic:
    global _client
    if _client is None:
        _client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    return _client


def _load(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def _pick_template(job: dict) -> Literal["lending", "risk"]:
    combined = ((job.get("title") or "") + " " + (job.get("description") or "")).lower()
    for signal in CONTROL_RISK_SIGNALS:
        if signal.lower() in combined:
            return "risk"
    return "lending"


# ---------------------------------------------------------------------------
# LibreOffice PDF conversion
# ---------------------------------------------------------------------------

def _docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert a Word document (bytes) to PDF via LibreOffice headless.

    LibreOffice is installed on Streamlit Cloud via packages.txt.
    fonts-crosextra-carlito provides Calibri-compatible metrics so the
    rendered PDF is visually identical to the Word document.

    Raises RuntimeError if LibreOffice is not found or conversion fails.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "doc.docx")
        pdf_path  = os.path.join(tmpdir, "doc.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                "libreoffice", "--headless", "--norestore",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if not os.path.exists(pdf_path):
            raise RuntimeError(
                f"LibreOffice PDF conversion failed.\n"
                f"stdout: {result.stdout}\nstderr: {result.stderr}"
            )

        with open(pdf_path, "rb") as f:
            return f.read()


# ---------------------------------------------------------------------------
# Haiku — tailor resume content
# ---------------------------------------------------------------------------

_TAILOR_PROMPT = """\
You are an expert ATS resume writer. Tailor Yang Yang's resume to maximise its score
in Applicant Tracking Systems (ATS) for the target role below.

You must work ONLY with experience already in the base resume — never invent anything.
The goal is to secure a first interview by scoring as high as possible against the ATS,
while remaining completely truthful.

=== ATS RULES — follow every one ===

KEYWORD STRATEGY
- Read the job description carefully. Identify every skill, tool, and phrase that
  appears 2 or more times — these are the highest-weight ATS keywords.
- Use those EXACT phrases (copy/paste verbatim, do not paraphrase) in the Summary,
  Skills list, and experience bullets wherever they accurately describe what Yang did.
  ATS matches exact strings. "controls assurance" ≠ "control assurance review".
- Include both spelled-out and abbreviated forms where relevant (e.g. "APRA" and
  "Australian Prudential Regulation Authority" if the JD uses both).

PROFESSIONAL SUMMARY
- Open sentence must mirror the target job title or role type as closely as possible.
  Example: if the role is "Senior Analyst – Controls Assurance", start with
  "Senior controls assurance professional with X years..."
- Pack the paragraph with high-frequency JD keywords. Every sentence should contain
  at least one exact JD phrase.
- Exactly 1 paragraph, 3-4 sentences. No more.

CORE SKILLS
- Include every skill term from the JD requirements that Yang genuinely has.
  If the JD says "risk in change" — list it exactly as "Risk in Change".
- 10-12 skills maximum. One skill per line, plain phrase, no bullets or dashes.
- Order by relevance: most JD-critical skills first.

EXPERIENCE BULLETS
- Lead every bullet with a strong action verb: Led, Conducted, Assessed, Identified,
  Delivered, Drove, Reviewed, Escalated, Monitored, Evaluated, Managed, Supported.
  Avoid weak openers: "Responsible for", "Assisted with", "Involved in".
- Use JD language verbatim in bullets where it accurately describes Yang's work.
- For the most recent role (Credit Risk Manager): write 8-10 bullets, surface the
  most JD-relevant work first.
- For ARM and Teller: 3-4 bullets each, focused on transferable relevance.
- Keep bullets factual. Never invent numbers, outcomes, or tools not in the base resume.

FIXED RULES
- Keep all employers, titles, dates, and education exactly as written.
- Never invent responsibilities, certifications, or outcomes.
- Output ONLY the structured text below — no commentary, no preamble.

=== OUTPUT FORMAT (exact) ===

SUMMARY
[single paragraph]

SKILLS
[Skill 1]
[Skill 2]
[...up to 12, one per line]

ROLE: Bank of China Australia – Credit Risk Manager
[bullet 1]
[bullet 2]
[...8-10 bullets]

ROLE: Bank of China Australia – Assistant Relationship Manager
[bullet 1]
[...3-4 bullets]

ROLE: Bank of China Australia – Teller
[bullet 1]
[...3-4 bullets]

Candidate profile: {profile}

Job title: {title}
Company: {company}
Job description: {description}

Base resume:
{resume}"""


def _tailor_content(job: dict, template: str) -> dict:
    """Call Haiku to tailor resume content. Returns parsed sections dict."""
    client = _get_client()
    profile = _load(_PROFILE_PATH)
    resume  = _load(_RESUME_RISK_PATH if template == "risk" else _RESUME_LENDING_PATH)

    resp = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=3000,
        messages=[{
            "role": "user",
            "content": _TAILOR_PROMPT.format(
                profile=profile[:2500],
                title=job.get("title", ""),
                company=job.get("company", ""),
                description=(job.get("description") or "")[:2000],
                resume=resume,
            ),
        }],
    )
    return _parse_tailored(resp.content[0].text.strip())


def _parse_tailored(text: str) -> dict:
    """Parse Haiku's structured output into sections."""
    sections = {"summary": [], "skills": [], "roles": {}}
    current = None
    current_role = None

    for line in text.splitlines():
        stripped = line.strip()
        if stripped == "SUMMARY":
            current = "summary"
        elif stripped == "SKILLS":
            current = "skills"
        elif stripped.startswith("ROLE:"):
            current = "role"
            current_role = stripped[5:].strip()
            sections["roles"][current_role] = []
        elif current == "summary" and stripped:
            sections["summary"].append(stripped)
        elif current == "skills" and stripped:
            skill = re.sub(r'^[-•–]\s*', '', stripped)
            if skill:
                sections["skills"].append(skill)
        elif current == "role" and current_role and stripped:
            bullet = re.sub(r'^[-•–]\s*', '', stripped)
            # Skip lines that look like location/date metadata e.g. "Sydney | 2019 – 2023"
            if bullet and not re.search(r'\|\s*\d{4}', bullet):
                sections["roles"][current_role].append(bullet)

    return sections


# ---------------------------------------------------------------------------
# Role-bullet matching helper
# ---------------------------------------------------------------------------

def _find_role_bullets(role_name: str, roles_dict: dict) -> list:
    """Return bullets for role_name from the parsed roles dict.

    Matches on the job-title portion after '–' so that roles sharing the
    same employer prefix ('Bank of China Australia') don't all resolve to
    the same entry.  Falls back to a full-name substring match.
    """
    def title_part(name: str) -> str:
        parts = name.split("–")
        return parts[-1].strip().lower() if len(parts) > 1 else name.lower()

    target = title_part(role_name)
    for key, items in roles_dict.items():
        if target[:18] in title_part(key) or title_part(key)[:18] in target:
            return items
    for key, items in roles_dict.items():
        if role_name.lower()[:30] in key.lower():
            return items
    return []


# ---------------------------------------------------------------------------
# Shared Word document helpers
# ---------------------------------------------------------------------------

def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin full-width horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, size_pt: float, bold: bool = False,
              color: RGBColor = BLACK, name: str = "Calibri") -> None:
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.color.rgb = color


def _section_heading(doc: Document, text: str, template: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    color = BLUE if template == "risk" else BLACK
    _set_font(run, 10.5, bold=True, color=color)


# ---------------------------------------------------------------------------
# Resume — Word document
# ---------------------------------------------------------------------------

def build_resume_docx(job: dict) -> bytes:
    """Build a tailored, ATS-optimised resume Word document. Returns bytes."""
    template = _pick_template(job)
    content  = _tailor_content(job, template)

    doc = Document()

    # Tighter margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)

    # --- Name ---
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(1)
    name_run = name_para.add_run("Yang Yang")
    _set_font(name_run, 18, bold=True)

    if template == "risk":
        _add_horizontal_rule(doc)

    # --- Contact (single line — ATS parsers handle | delimiters fine) ---
    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(5)
    c = contact_para.add_run(
        "Sydney NSW  |  Australian Permanent Resident  |  0401 877 625  |  yy.lu.33@gmail.com"
    )
    _set_font(c, 9.5, color=DARK_GREY)

    if template == "lending":
        _add_horizontal_rule(doc)

    # --- Professional Summary ---
    _section_heading(doc, "Professional Summary", template)
    summary_text = " ".join(content["summary"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(summary_text)
    _set_font(run, 10.5)

    if template == "lending":
        _add_horizontal_rule(doc)

    # --- Core Skills (one per line — ATS extracts skills line-by-line) ---
    _section_heading(doc, "Core Skills", template)
    sk_prefix = "–" if template == "risk" else "•"
    for skill in content["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)
        run = p.add_run(f"{sk_prefix}  {skill}")
        _set_font(run, 10.5)

    if template == "lending":
        _add_horizontal_rule(doc)

    # --- Professional Experience ---
    _section_heading(doc, "Professional Experience", template)

    roles_order = [
        ("Bank of China Australia – Credit Risk Manager",
         "Sydney Head Office  |  2019 – 2023"),
        ("Bank of China Australia – Assistant Relationship Manager",
         "Hurstville Branch  |  2016 – 2019"),
        ("Bank of China Australia – Teller",
         "Hurstville Branch  |  2013 – 2015"),
    ]

    for role_name, role_meta in roles_order:
        # Role title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(role_name)
        _set_font(run, 10.5, bold=True)

        # Location | dates
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(role_meta)
        _set_font(run2, 9.5, color=DARK_GREY)

        # Bullets — hanging indent so wrapped lines clear the prefix character
        bullets = _find_role_bullets(role_name, content["roles"])
        for bullet_text in bullets:
            p3 = doc.add_paragraph(
                style='List Bullet' if template == "lending" else 'Normal'
            )
            p3.paragraph_format.space_after       = Pt(1)
            p3.paragraph_format.left_indent        = Inches(0.2)
            p3.paragraph_format.first_line_indent  = Inches(-0.2)
            prefix_char = "–" if template == "risk" else "•"
            run3 = p3.add_run(f"{prefix_char}  {bullet_text}")
            _set_font(run3, 10.5)

        if template == "lending":
            _add_horizontal_rule(doc)

    # --- Education ---
    _section_heading(doc, "Education", template)
    for line in [
        "Master of Finance – Western Sydney University",
        "Bachelor of Accounting – Southern Cross University",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        _set_font(run, 10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_resume_pdf(job: dict) -> bytes:
    """Generate a PDF resume that is pixel-identical to the Word version.

    Builds the .docx first, then converts via LibreOffice headless so both
    files share exactly the same layout, fonts, and spacing.
    """
    return _docx_bytes_to_pdf(build_resume_docx(job))


# ---------------------------------------------------------------------------
# Cover letter — Word document
# ---------------------------------------------------------------------------

def build_cover_letter_docx(text: str) -> bytes:
    """Wrap cover letter text in a clean Word document."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)

    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_cover_letter_pdf(text: str) -> bytes:
    """Generate a PDF cover letter identical to the Word version via LibreOffice."""
    return _docx_bytes_to_pdf(build_cover_letter_docx(text))
